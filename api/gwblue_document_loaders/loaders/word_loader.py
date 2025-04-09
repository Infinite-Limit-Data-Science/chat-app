from typing import Union, Iterator
from pathlib import Path
import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.section import Section
from lxml import etree

from langchain_core.documents import Document
from .base_loader import BaseLoader

EMU_PER_POINT = 12700.0

class WordLoader(BaseLoader):
    """
    Recursively extract text from paragraphs, tables, headers, and footers,
    splitting on explicit page breaks <w:br w:type="page" />.
    Each 'page' of text is returned as a Document, with 'page_number' metadata.
    """
    def __init__(self, file_path: Union[str, Path]):
        self._file_path = file_path
        self.doc = docx.Document(self._file_path)
        self._mean_font_size = self._compute_mean_font_size()
        self._counters = {}

    def lazy_load(self) -> Iterator[Document]:
        page_lines = []
        page_number = 1

        def yield_page():
            nonlocal page_number
            if page_lines:
                text = "\n".join(page_lines).strip()
                if text:
                    metadata = {
                        "source": self.sourcify(self._file_path),
                        #"page_number": page_number,
                    }
                    yield Document(page_content=text, metadata=metadata)
                page_lines.clear()
                page_number += 1

        for element in self.doc.element.body:
            if element.tag.endswith("p"):
                paragraph_obj = docx.text.paragraph.Paragraph(element, self.doc)

                paragraph_text_fragments = self._split_paragraph_by_page_break(paragraph_obj)

                for frag in paragraph_text_fragments:
                    if frag == "__PAGE_BREAK__":
                        yield from yield_page()
                    else:
                        final_line = self._apply_bullet_and_importance(paragraph_obj, frag)
                        page_lines.append(final_line)

            elif element.tag.endswith("tbl"):
                table = docx.table.Table(element, self.doc)
                table_str = self.table_to_markdown(table)
                if table_str:
                    page_lines.append(table_str)

        yield from yield_page()

    load = lazy_load

    def _split_paragraph_by_page_break(self, paragraph: Paragraph) -> list[str]:
        fragments = []
        current_chunk = []

        for run in paragraph.runs:
            run_elm = run._r
            br_elems = run_elm.xpath(
                ".//*[local-name()='br' and @*[local-name()='type']='page']"
            )
            if br_elems:
                before_br = run.text
                if before_br:
                    current_chunk.append(before_br)

                if current_chunk:
                    fragments.append("".join(current_chunk).strip())
                    current_chunk = []

                fragments.append("__PAGE_BREAK__")
            else:
                current_chunk.append(run.text)

        if current_chunk:
            fragments.append("".join(current_chunk).strip())

        return [frag for frag in fragments if frag]

    def _apply_bullet_and_importance(self, paragraph: Paragraph, raw_text: str) -> str:
        text_with_bullet = self._apply_bullet_numbering(paragraph, raw_text)
        if self._mean_font_size > 0:
            if self._is_paragraph_header(paragraph, self._mean_font_size, factor=1.3):
                return f"<important title='{text_with_bullet}'>{text_with_bullet}</important>"
        return text_with_bullet

    def _compute_mean_font_size(self) -> float:
        sizes_in_points = []
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.font.size is not None:
                    sizes_in_points.append(run.font.size / EMU_PER_POINT)
        if not sizes_in_points:
            return 0.0
        return sum(sizes_in_points) / len(sizes_in_points)

    def _is_paragraph_header(self, paragraph: Paragraph, baseline: float, factor: float = 1.3) -> bool:
        largest_run_size = 0.0
        for run in paragraph.runs:
            if run.font.size is not None:
                size_pt = run.font.size / EMU_PER_POINT
                if size_pt > largest_run_size:
                    largest_run_size = size_pt
        return largest_run_size >= (factor * baseline)

    def _apply_bullet_numbering(self, paragraph: Paragraph, raw_text: str) -> str:
        p = paragraph._p
        if p is None or p.pPr is None or p.pPr.numPr is None:
            return raw_text

        numId_elm = p.pPr.numPr.numId
        ilvl_elm  = p.pPr.numPr.ilvl
        if numId_elm is None or ilvl_elm is None:
            return raw_text

        try:
            numId = int(numId_elm.val)
            ilvl  = int(ilvl_elm.val)
        except (TypeError, ValueError):
            return raw_text

        num_fmt, lvl_text = self._get_list_format(numId, ilvl)
        if not num_fmt:
            return raw_text

        if num_fmt == "bullet":
            prefix = lvl_text
        elif num_fmt in ("decimal", "upperRoman", "lowerRoman", "upperLetter", "lowerLetter"):
            key = (numId, ilvl)
            if key not in self._counters:
                self._counters[key] = 0
            self._counters[key] += 1
            prefix = f"{self._counters[key]}."
        else:
            prefix = lvl_text.replace("%", "").strip()

        return f"{prefix} {raw_text}"

    def _get_list_format(self, numId: int, ilvl: int) -> tuple[str, str]:
        numbering_part = self.doc.part.numbering_part
        numbering_el = numbering_part.element
        numbering_xml_str = numbering_el.xml
        root = etree.fromstring(numbering_xml_str.encode("utf-8"))
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        num_elems = root.xpath(f".//w:num[@w:numId='{numId}']", namespaces=nsmap)
        if not num_elems:
            return ("", "")

        abstract_refs = num_elems[0].xpath("./w:abstractNumId/@w:val", namespaces=nsmap)
        if not abstract_refs:
            return ("", "")

        abstract_id = abstract_refs[0]
        abs_elems = root.xpath(f".//w:abstractNum[@w:abstractNumId='{abstract_id}']", namespaces=nsmap)
        if not abs_elems:
            return ("", "")

        lvl_elems = abs_elems[0].xpath(f"./w:lvl[@w:ilvl='{ilvl}']", namespaces=nsmap)
        if not lvl_elems:
            return ("", "")

        num_fmt_vals = lvl_elems[0].xpath("./w:numFmt/@w:val", namespaces=nsmap)
        lvl_text_vals = lvl_elems[0].xpath("./w:lvlText/@w:val", namespaces=nsmap)
        if not num_fmt_vals or not lvl_text_vals:
            return ("", "")

        return (num_fmt_vals[0], lvl_text_vals[0])

    def table_to_markdown(self, table: Table) -> str:
        if not table.rows:
            return ""
        
        n_cols = len(table.rows[0].cells)

        all_rows = []
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                nested_tables_md = []
                if cell.tables:
                    for nested_table in cell.tables:
                        nested_tables_md.append(self.table_to_markdown(nested_table))
                
                cell_text = cell.text.strip()
                if nested_tables_md:
                    nested_md_str = "\n\n".join(nested_tables_md)
                    cell_content = f"{cell_text}\n\n{nested_md_str}"
                else:
                    cell_content = cell_text
                
                row_data.append(cell_content)
            all_rows.append(row_data)

        if len(all_rows) > 1:
            header_cols = all_rows[0]
            while len(header_cols) < n_cols:
                header_cols.append("")
            header_line = "| " + " | ".join(header_cols) + " |"

            delimiter = "| " + " | ".join(["---"] * n_cols) + " |"

            body_lines = []
            for row in all_rows[1:]:
                while len(row) < n_cols:
                    row.append("")                
                line = "| " + " | ".join(row) + " |"
                body_lines.append(line)

            md_table = "\n".join([header_line, delimiter] + body_lines)
        else:
            row = all_rows[0]
            while len(row) < n_cols:
                row.append("")
            md_table = "| " + " | ".join(row) + " |\n"
            md_table += "| " + " | ".join(["---"] * n_cols) + " |\n"

        return md_table

    def extract_header_footer_text(self, section: Section) -> str:
        parts = []
        if section.header:
            for paragraph in section.header.paragraphs:
                p_text = paragraph.text.strip()
                if p_text:
                    parts.append(p_text)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                p_text = paragraph.text.strip()
                if p_text:
                    parts.append(p_text)
        return "\n".join(parts)