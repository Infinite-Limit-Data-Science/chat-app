from typing import Union, Iterator
from pathlib import Path
import docx
from docx.table import Table
from docx.section import Section
from langchain_core.documents import Document
from orchestrators.doc.document_loaders.base_loader import BaseLoader

class WordLoader(BaseLoader):
    """
    Recursively extract all text from all paragraphs, tables, headers, and footers of document

    Note table cells can have tables and paragraphs
    The text property of a cell only extracts text from the cell's paragraphs and not tables
    Hence, recursion is required for tables within cells as well
    """
    def __init__(self, file_path: Union[str, Path]):
        self._file_path = file_path
        self.doc = docx.Document(file_path)

    def lazy_load(self) -> Iterator[Document]:
        """Lazily load document"""
        full_text = []

        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                paragraph = docx.text.paragraph.Paragraph(element, self.doc)
                full_text.append(paragraph.text)
            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, self.doc)
                full_text.append(self.extract_table_text(table))

        for section in self.doc.sections:
            full_text.append(self.extract_header_footer_text(section))

        metadata = {'source': str(self._file_path)}
        yield Document(page_content="\n".join(full_text), metadata=metadata)
    load = lazy_load

    def extract_table_text(self, table: Table) -> str:
        """Recursively extracts text from tables, including nested tables."""
        table_text = []
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text.strip())
                if cell.tables:
                    for nested_table in cell.tables:
                        table_text.append(self.extract_table_text(nested_table))

        return "\n".join(table_text)

    def extract_header_footer_text(self, section: Section) -> str:
        header_footer_text = []

        if section.header:
            for paragraph in section.header.paragraphs:
                header_footer_text.append(paragraph.text)

        if section.footer:
            for paragraph in section.footer.paragraphs:
                header_footer_text.append(paragraph.text)

        return "\n".join(header_footer_text)